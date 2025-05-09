from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from PIL import Image
import io
from docx2pdf import convert

# Function for natural sorting (handles numeric parts of filenames correctly)
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

# Function to add a title heading with capital letters
def add_title(doc, title):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title.upper())  # Convert title to uppercase
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 255, 0)  # Green color
    doc.add_paragraph()  # Add a blank line after the title

# Function to add a 3D border to a page (section)
def add_3d_border(section):
    # Get the section properties element
    sectPr = section._sectPr
    
    # Create page borders element
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    # Set 3D border style values for all four sides
    for border_side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_side}')
        border.set(qn('w:val'), 'threeDEngrave')  # 3D engraved style
        border.set(qn('w:sz'), '24')  # 3 points (24 eighths of a point)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        pgBorders.append(border)
    
    # Add the page borders element to the section properties
    sectPr.append(pgBorders)

# Function to crop an image from the bottom by a specified amount
def crop_image_from_bottom(image_path, crop_cm=0.2):
    try:
        # Open the image
        with Image.open(image_path) as img:
            # Calculate crop amount in pixels based on resolution
            # Assume 96 DPI (standard screen resolution)
            # 1 cm = 96/2.54 pixels (approximately 37.8 pixels)
            dpi = 96
            crop_pixels = int((crop_cm * dpi) / 2.54)
            
            # Calculate new height
            new_height = img.height - crop_pixels
            if new_height <= 0:
                print(f"Warning: Crop amount too large for {image_path}. Using original image.")
                return image_path
            
            # Crop the image
            cropped_img = img.crop((0, 0, img.width, new_height))
            
            # Create a BytesIO object to hold the cropped image
            img_byte_arr = io.BytesIO()
            cropped_img.save(img_byte_arr, format=img.format)
            img_byte_arr.seek(0)
            
            return img_byte_arr
    except Exception as e:
        print(f"Error cropping image {image_path}: {e}")
        return image_path  # Return original path if cropping fails

# Main folder containing subfolders (lectures)
main_folder = r"G:\work\reasoning practice"  # Update this path

# Iterate through each subfolder (lecture)
for subfolder_name in os.listdir(main_folder):
    subfolder_path = os.path.join(main_folder, subfolder_name)
    
    # Check if it's a directory
    if os.path.isdir(subfolder_path):
        # Create a new Word document
        doc = Document()
        
        # Set page margins (left, right, top, bottom)
        section = doc.sections[0]
        section.left_margin = Cm(1.0)  # 1.0 cm left margin
        section.right_margin = Cm(1.0)  # 1.0 cm right margin
        section.top_margin = Cm(1.0)  # 1.0 cm top margin
        section.bottom_margin = Cm(1.0)  # 1.0 cm bottom margin
        
        # Add 3D border to the section
        add_3d_border(section)
        
        # Add title heading to the first page
        add_title(doc, subfolder_name)
        
        # Get a list of PNG files in the subfolder
        png_files = [f for f in os.listdir(subfolder_path) if f.endswith('.png')]
        
        # Sort files naturally (handles numeric parts correctly)
        png_files.sort(key=natural_sort_key)
        
        # Print the sorted file list for debugging
        print(f"Files in {subfolder_name} will be processed in this order:")
        for f in png_files:
            print(f"  {f}")
        
        # Track how many images have been added to the current page
        images_on_current_page = 0
        current_page_number = 1
        
        # Insert PNGs into the Word document
        for i, png_file in enumerate(png_files):
            img_path = os.path.join(subfolder_path, png_file)
            
            # Crop the image from the bottom by 0.2 cm
            cropped_img = crop_image_from_bottom(img_path, crop_cm=0.2)
            
            # Determine how many images should be on this page
            images_per_page = 2 if current_page_number == 1 else 3
            
            # Set image dimensions based on the page
            if current_page_number == 1:
                # First page: 2 images with dimensions 9.05 cm x 18.46 cm
                img_width = Cm(18.46)
                img_height = Cm(9.05)
            else:
                # Subsequent pages: 3 images with dimensions 7.68 cm x 16.41 cm
                img_width = Cm(16.41)
                img_height = Cm(7.68)
            
            # Add the image to the document with CENTER alignment
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the paragraph containing the image
            run = paragraph.add_run()
            
            # Add the cropped image to the document
            run.add_picture(cropped_img, width=img_width, height=img_height)
            
            # Increment counter for images on current page
            images_on_current_page += 1
            
            # Check if we need to add a page break
            is_last_image = (i == len(png_files) - 1)
            
            if not is_last_image and images_on_current_page >= images_per_page:
                # Add page break
                doc.add_page_break()
                
                # Add 3D border to the new page's section
                new_section = doc.add_section()
                new_section.left_margin = Cm(1.0)
                new_section.right_margin = Cm(1.0)
                new_section.top_margin = Cm(1.0)
                new_section.bottom_margin = Cm(1.0)
                add_3d_border(new_section)
                
                # Reset counter and increment page number
                images_on_current_page = 0
                current_page_number += 1
        
        # Save the Word document with the subfolder name
        try:
            doc.save(f"{subfolder_name}.docx")
            print(f"Created: {subfolder_name}.docx")
        except PermissionError:
            print(f"Error: Permission denied for {subfolder_name}.docx. Please close the file or check permissions.")
            continue
